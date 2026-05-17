import io
import pytest
from httpx import AsyncClient, ASGITransport
import pypdf
import docx
import openpyxl
from app.main import app
from app.api.deps import get_current_active_user
from app.models.user import User

def create_sample_pdf() -> io.BytesIO:
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=72*8.5, height=72*11)
    # pypdf doesn't easily create text pages from scratch without reportlab,
    # but let's test if parser handles empty page or let's mock it
    pdf_stream = io.BytesIO()
    writer.write(pdf_stream)
    pdf_stream.seek(0)
    return pdf_stream

def create_sample_docx() -> io.BytesIO:
    doc = docx.Document()
    doc.add_heading('Sample Title', 0)
    doc.add_paragraph('This is a test paragraph.')
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Val 1'
    table.rows[1].cells[1].text = 'Val 2'
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

def create_sample_excel() -> io.BytesIO:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TestSheet"
    ws.append(["ColA", "ColB"])
    ws.append(["Data1", "Data2"])
    excel_stream = io.BytesIO()
    wb.save(excel_stream)
    excel_stream.seek(0)
    return excel_stream

@pytest.mark.asyncio
async def test_parse_endpoints():
    mock_user = User(id=1, username="test", role="admin", status="active")
    app.dependency_overrides[get_current_active_user] = lambda: mock_user
    
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        # Test DOCX
        docx_bytes = create_sample_docx().read()
        files = {"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = await ac.post("/api/v1/parsers/upload", files=files)
        assert res.status_code == 200
        data = res.json()["data"]
        assert data["file_type"] == "docx"
        assert "This is a test paragraph." in data["content"]
        assert len(data["chunks"]) >= 1

        # Test Excel
        excel_bytes = create_sample_excel().read()
        files = {"file": ("test.xlsx", excel_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        res = await ac.post("/api/v1/parsers/upload", files=files)
        assert res.status_code == 200
        data = res.json()["data"]
        assert data["file_type"] == "xlsx"
        assert "Data1" in data["content"]
        
        # Test PDF
        pdf_bytes = create_sample_pdf().read()
        files = {"file": ("test.pdf", pdf_bytes, "application/pdf")}
        res = await ac.post("/api/v1/parsers/upload", files=files)
        assert res.status_code == 200
        data = res.json()["data"]
        assert data["file_type"] == "pdf"
        
        # Test Unsupported
        files = {"file": ("test.txt", b"hello", "text/plain")}
        res = await ac.post("/api/v1/parsers/upload", files=files)
        assert res.status_code == 400
