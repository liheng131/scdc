"""
HTML → Word 转换器 (v2)

v2: 从 HTML slide 结构提取内容 + 保留图片、表格、格式
v1: BeautifulSoup 纯文本提取（丢失所有样式和图片）

流程：
1. 解析 <section class="slide"> 块
2. 每个 slide → 标题（Heading 1）+ 段落 + 图片 + 表格
3. 图片从 base64 解码后内嵌到 Word（不是链接）
4. slide 之间用分页符
"""

import base64
import io
import logging
import os
import tempfile
from typing import Optional

import docx
import docx.shared
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class HTMLToDOCXConverter:
    """HTML → Word 转换器 (v2)"""

    async def convert(self, html: str, output_path: str):
        """将 HTML 转换为 Word 文档

        Args:
            html: 完整的 HTML 字符串
            output_path: Word 输出路径
        """
        logger.info("Converting HTML to Word (v2: with embedded images)...")

        soup = BeautifulSoup(html, "html.parser")
        doc = docx.Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = docx.shared.Pt(12)

        # 遍历所有 slide
        slides = soup.find_all("section", class_="slide")
        if not slides:
            logger.warning("No .slide elements found, using body content")
            slides = [soup.body] if soup.body else [soup]

        for i, slide in enumerate(slides):
            logger.debug(f"Processing slide {i + 1}/{len(slides)}")

            # 提取并跳过 notes 块
            notes = slide.find("div", class_="notes")
            if notes:
                notes.extract()

            # 标题 h1/h2
            title_tag = slide.find(["h1", "h2"])
            if title_tag:
                title_text = title_tag.get_text(strip=True)
                if title_text:
                    doc.add_heading(title_text, level=1)

            # Kicker
            kicker_tag = slide.find("p", class_="kicker")
            if kicker_tag:
                kicker_text = kicker_tag.get_text(strip=True)
                if kicker_text:
                    p = doc.add_paragraph(kicker_text)
                    run = p.runs[0]
                    run.font.color.rgb = docx.shared.RGBColor(0x3B, 0x6C, 0xFF)
                    run.font.size = docx.shared.Pt(11)
                    run.bold = True

            # 段落
            for p_tag in slide.find_all("p"):
                cls = p_tag.get("class", [])
                if "kicker" in cls:
                    continue  # 已在上面处理
                if p_tag.find_parent("div", class_="notes"):
                    continue  # 跳过 notes

                text = p_tag.get_text(strip=True)
                if not text:
                    continue

                para = doc.add_paragraph()
                run = para.add_run(text)

                # 样式
                if "lede" in cls:
                    run.font.size = docx.shared.Pt(16)
                    run.font.color.rgb = docx.shared.RGBColor(0x55, 0x59, 0x6A)
                elif "dim" in cls or "dim2" in cls:
                    run.font.size = docx.shared.Pt(10)
                    run.font.color.rgb = docx.shared.RGBColor(0x8A, 0x8F, 0x9E)
                else:
                    run.font.size = docx.shared.Pt(12)

            # 图片（从 <img> 标签提取）
            for img_tag in slide.find_all("img"):
                src = img_tag.get("src", "")
                alt = img_tag.get("alt", "")
                if src.startswith("data:image"):
                    await self._embed_base64_image(doc, src, alt)
                elif src.startswith("http"):
                    # 网络图片 → 插入为占位符引用
                    logger.debug(f"Skipping remote image: {src[:80]}")
                    caption_text = alt or f"[图片: {src[:60]}...]"
                    p = doc.add_paragraph()
                    run = p.add_run(caption_text)
                    run.font.size = docx.shared.Pt(10)
                    run.font.italic = True

            # 表格
            for table_tag in slide.find_all("table"):
                self._add_table(doc, table_tag)

            # 页面分隔（除了最后一页）
            if i < len(slides) - 1:
                doc.add_page_break()

        doc.save(output_path)
        logger.info(f"Word document saved: {output_path}")

    async def _embed_base64_image(self, doc, data_url: str, alt: str = ""):
        """将 base64 图片解码并内嵌到 Word 文档"""
        try:
            header, encoded = data_url.split(",", 1)
            img_bytes = base64.b64decode(encoded)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name

            try:
                # 限制图片最大宽度为 6 英寸
                doc.add_picture(tmp_path, width=docx.shared.Inches(5.5))
                if alt:
                    p = doc.add_paragraph()
                    run = p.add_run(alt)
                    run.font.size = docx.shared.Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = docx.shared.RGBColor(0x8A, 0x8F, 0x9E)
            finally:
                os.unlink(tmp_path)
        except Exception as e:
            logger.warning(f"Failed to embed base64 image in DOCX: {e}")

    def _add_table(self, doc, html_table):
        """将 HTML 表格转换为 Word 表格"""
        rows = html_table.find_all("tr")
        if not rows:
            return

        # 确定列数
        first_row_cells = rows[0].find_all(["th", "td"])
        num_cols = len(first_row_cells)
        if num_cols == 0:
            return

        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for j, cell in enumerate(cells):
                if j >= num_cols:
                    break
                cell_text = cell.get_text(strip=True)
                table.rows[i].cells[j].text = cell_text

                # 表头样式
                if i == 0 and row.find("th"):
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = docx.shared.Pt(11)
