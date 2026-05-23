"""
Word 文档解析器

使用 python-docx 库解析 .docx 文件，提取段落文本和表格内容。

为什么分别处理段落和表格：
- 段落（paragraph）和表格（table）在 Word 文档中是独立的元素类型
- 表格中的信息（如数据对比、统计表格）通常具有独立的语义价值
- metadata 中标记 chunk 类型（paragraph/table）便于 ReporterAgent 区分展示格式
"""

import io
from typing import BinaryIO
import docx
from app.parsers.base import BaseParser
from app.schemas.parser import ParseResult, Chunk
from app.core.exceptions import BusinessException

class DocxParser(BaseParser):
    async def parse(self, file_stream: BinaryIO, filename: str) -> ParseResult:
        try:
            doc = docx.Document(file_stream)
            full_text = []
            chunks = []
            chunk_idx = 1

            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    full_text.append(text)
                    chunks.append(Chunk(
                        index=chunk_idx,
                        content=text,
                        metadata={"type": "paragraph"}
                    ))
                    chunk_idx += 1

            for t_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_text.append(row_text)
                if table_text:
                    content = "\n".join(table_text)
                    full_text.append(content)
                    chunks.append(Chunk(
                        index=chunk_idx,
                        content=content,
                        metadata={"type": "table", "table_index": t_idx + 1}
                    ))
                    chunk_idx += 1

            content = "\n\n".join(full_text)
            return ParseResult(
                filename=filename,
                file_type="docx",
                content=content,
                chunks=chunks,
                metadata={"total_paragraphs": len(doc.paragraphs), "total_tables": len(doc.tables)}
            )
        except Exception as e:
            raise BusinessException(code=422, message=f"Failed to parse DOCX file: {str(e)}")
